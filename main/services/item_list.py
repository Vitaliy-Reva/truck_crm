from docx import Document
from django.db.models.query import QuerySet

def get_items(items):
    document = Document()

    if type(items) == QuerySet:
        for item in items:
            document.add_paragraph(f"ID: {item.id}")
            document.add_paragraph(f"Водій: {item.driver_id.first_name} {item.driver_id.last_name}")
            document.add_paragraph(f"Транспорт: {item.transport_id.transport_model} ({item.transport_id.license_plate})")
            document.add_paragraph(f"Клієнт: {item.client_id.client_type} {item.client_id.first_name} {item.client_id.last_name}")
            document.add_paragraph(f"Замовлення: {item.order_id.order_name}")
            document.add_paragraph(f"Точка виїзду: {item.start_point}")
            document.add_paragraph(f"Точка прибуття: {item.end_point}")
            document.add_paragraph(f"Статус доставки: {item.status}")
            document.add_paragraph(f"Дистанція: {item.distance}")
            document.add_paragraph(f"Надвитрата пального: {item.fuel_status}")
            document.add_paragraph(f"Реальна кількість пального: {item.fuel_actual}")
            document.add_paragraph(f"Заплановано пального (км): {item.fuel_planned}")
            document.add_paragraph("")

        document.save('test_list.docx')
    
    else:
            document.add_paragraph(f"ID: {items.id}")
            document.add_paragraph(f"Водій: {items.driver_id.first_name} {items.driver_id.last_name}")
            document.add_paragraph(f"Транспорт: {items.transport_id.transport_model} ({items.transport_id.license_plate})")
            document.add_paragraph(f"Клієнт: {items.client_id.client_type} {items.client_id.first_name} {items.client_id.last_name}")
            document.add_paragraph(f"Замовлення: {items.order_id.order_name}")
            document.add_paragraph(f"Точка виїзду: {items.start_point}")
            document.add_paragraph(f"Точка прибуття: {items.end_point}")
            document.add_paragraph(f"Статус доставки: {items.status}")
            document.add_paragraph(f"Дистанція: {items.distance}")
            document.add_paragraph(f"Надвитрата пального: {items.fuel_status}")
            document.add_paragraph(f"Реальна кількість пального: {items.fuel_actual}")
            document.add_paragraph(f"Заплановано пального (км): {items.fuel_planned}")
        
    document.save('test_list.docx')